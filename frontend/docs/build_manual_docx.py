"""Convert USER_MANUAL_DE.md to USER_MANUAL_DE.docx"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DOCS = Path(__file__).resolve().parent
MD_FILE = DOCS / "USER_MANUAL_DE.md"
OUT_FILE = DOCS / "USER_MANUAL_DE.docx"


def add_formatted_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)
    p.paragraph_format.space_after = Pt(6)
    return p


def parse_table_row(line):
    line = line.strip()
    if not line.startswith("|"):
        return None
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def is_separator_row(cells):
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells)


def md_to_docx(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = parse_table_row(tl)
                if cells and not is_separator_row(cells):
                    rows.append(cells)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row_cells in enumerate(rows):
                    for ci, cell_text in enumerate(row_cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                        run = p.add_run(clean)
                        if ri == 0:
                            run.bold = True
                doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                item = lines[i].strip()[2:]
                add_formatted_paragraph(doc, item, style="List Bullet")
                i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                add_formatted_paragraph(doc, item, style="List Number")
                i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph(stripped.strip("*"))
            p.paragraph_format.space_before = Pt(12)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "---"
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or re.match(r"^\d+\.\s", nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        add_formatted_paragraph(doc, " ".join(para_lines))

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else MD_FILE
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT_FILE
    md_to_docx(md, out)
