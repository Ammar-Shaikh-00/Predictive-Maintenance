"""Generate concise ML_FRONTEND_PAYLOAD_CONTRACT.docx"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BACKEND = "http://192.168.100.24:8002"
OUT = Path(__file__).with_name("ML_FRONTEND_PAYLOAD_CONTRACT.docx")


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def p(text: str) -> None:
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    def table(headers: list[str], rows: list[list[str]]) -> None:
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1F4E79")
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = tbl.rows[r_i + 1].cells[c_i]
                cell.text = str(val)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
        doc.add_paragraph()

    h("ML → Frontend Contract (short)", 0)
    p("For: Frontend teammate | From: Ammar (live_monitor) | Date: 12 Aug 2026")
    p(f"Use backend only: {BACKEND}  |  Do not use localhost.")

    h("What live_monitor does", 1)
    p(
        "Every ~10s it polls the extruder, builds a 5-min window, detects machine "
        "state (ML), scores anomaly (ML), compares features to baselines, then "
        "writes results to Postgres via backend APIs."
    )
    p(
        "Your job: read those Postgres APIs and show the results in Operations "
        "Center. Do not invent Accuracy %."
    )

    h("APIs to read", 1)
    table(
        ["Endpoint", "What it is", "Purpose in UI"],
        [
            [
                "GET /live-process-windows",
                "One process window (features + confirmed state)",
                "Show current state and key live values",
            ],
            [
                "GET /live-run-evaluations",
                "Overall evaluation for that window",
                "Module 7 Live AI card + status (NORMAL/WARNING/CRITICAL)",
            ],
            [
                "GET /live-feature-evaluations",
                "Per-feature z-score vs baseline",
                "Why it is warning — heatmap / drivers",
            ],
            [
                "GET /baseline-registry",
                "Normal ranges per regime (LOW/MID/HIGH)",
                "Optional reference bands",
            ],
        ],
    )

    h("Main fields on live-run-evaluations", 1)
    table(
        ["Field", "Meaning", "UI use"],
        [
            [
                "detected_state",
                "Confirmed machine state (OFF, HEATING, PRODUCTION, …)",
                "State label",
            ],
            [
                "overall_status",
                "NORMAL / WARNING / CRITICAL",
                "Status chip / traffic light",
            ],
            [
                "explanation_text",
                "Plain-language summary + recommended action",
                "Module 7 AI recommendation text",
            ],
            [
                "ml_is_anomaly",
                "True if Isolation Forest flagged unusual behavior",
                "Anomaly badge",
            ],
            [
                "ml_anomaly_score",
                "Anomaly intensity 0–1",
                "Optional score",
            ],
            [
                "drift_score",
                "How far process drifted from baseline (0–1)",
                "Drift indicator",
            ],
            [
                "stability_status",
                "STABLE / TRANSITION / UNSTABLE",
                "Stability label",
            ],
            [
                "active_regime",
                "Pressure regime LOW / MID / HIGH",
                "Context",
            ],
        ],
    )

    h("How to wire Modules 7 / 15 / 16", 1)
    p(
        "Module 7 (Live AI): show latest explanation_text. Prefer this over demo "
        "text when present."
    )
    p(
        "Module 15 (Predictions/risks): for now, build cards from WARNING/CRITICAL "
        "rows in live-feature-evaluations, and/or ml_is_anomaly=true."
    )
    p(
        "Module 16 (Actions): use the Recommended action sentence inside "
        "explanation_text."
    )

    h("Quick rules", 1)
    p("1) Base URL = backend host above (same as rest of app).")
    p(
        "2) Poll latest: /live-run-evaluations?limit=1 "
        "(or filter by production_run_id)."
    )
    p("3) Never call live_monitor local ports from the product UI.")
    p(
        "4) Labels like [MODEL_PREDICTION] / [RULE_BASED] in explanation_text are "
        "provenance — show as badges if useful."
    )

    doc.save(OUT)
    print(f"saved {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
