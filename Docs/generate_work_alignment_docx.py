"""Generate WORK_ALIGNMENT_ARCHITECTURE_vs_TEAMS.docx"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, bold=False, size=11):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=11)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)


def shade_header(row):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, size=10)
    shade_header(hdr)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading(
        "Work Alignment — New Architecture vs Current Status", 0
    )
    for run in title.runs:
        set_run_font(run, bold=True, size=18)

    for line in [
        "Date: 11 Aug 2026",
        "Source of truth (product vision): Docs/NEW_ARCHITECTURE_EN.doc (+ German Docs/new architecture.docx)",
        "Repo: Predictive Maintenance / ZITTA Production Intelligence",
        "Storage rule: Postgres only (via backend APIs). SQLite is legacy and must not be used for new work.",
    ]:
        add_para(doc, line)

    doc.add_heading("1. Team ownership (clear split)", 1)
    add_table(
        doc,
        ["Area", "Owner", "Folder / surface"],
        [
            [
                "AI / ML pipeline",
                "Ammar",
                "live_monitor/ (models, features, state, anomaly, evaluation, retrain)",
            ],
            ["Backend APIs + Postgres", "Teammate", "backend/"],
            [
                "Frontend / Operations Center UI",
                "Teammate",
                "frontend/ (UI already complete per screenshots)",
            ],
        ],
    )
    add_para(
        doc,
        "Rule: Ammar produces ML results and writes them to backend APIs. "
        "Teammate stores them in Postgres and displays them in the finished UI "
        "with honest provenance labels.",
    )

    doc.add_heading("2. Architecture rules (must stay aligned)", 1)
    add_para(doc, "From the New Architecture document:")
    add_bullets(
        doc,
        [
            "Product is a Digital Production Control Room, not a generic IT dashboard.",
            "Never mix: Digitalization progress ≠ Data quality ≠ Model Accuracy.",
            "Until models are validated, UI must show Prediction Readiness / Willingness to Predict — not fake Accuracy.",
            "Every displayed value needs provenance: LIVE | RULE_BASED | DERIVED | SIMULATED | MODEL_PREDICTION | MANUAL.",
            "System can monitor well today; full scrap / RUL / quality prediction needs more data sources later.",
        ],
    )

    doc.add_heading("3. Module map — Done vs Remaining", 1)
    add_para(
        doc,
        "Legend: Done | Partial | Remaining | Blocked on data / other team",
    )

    doc.add_heading("3.1 Product modules (architecture modules 1–20)", 2)
    add_table(
        doc,
        ["#", "Module", "Status", "Who", "Notes"],
        [
            ["1", "Production overview (OC home)", "Done", "Frontend", "Screenshots show OC cockpit"],
            ["2", "Factory / line map", "Done", "Frontend", "Extruders, 1/20 connected, grey machines"],
            ["3", "Digitalization progress", "Done", "Frontend + Backend", "Checklist / missing sources"],
            ["4", "Prediction readiness", "Done", "Frontend + Backend", "Shown as readiness (~38%), not validated Accuracy"],
            ["5", "Accuracy center (real model metrics)", "Remaining", "AI/ML + Backend + Frontend", "Only after validated model_versions; do not invent %"],
            ["6", "Live production status cards", "Done / Partial", "Frontend + live sensors", "Live trends exist; energy/throughput/scrap still missing sources"],
            ["7", "Live AI analysis (plain language)", "Partial", "AI/ML produces → Frontend displays", "ML scores exist in Postgres; OC still mostly demo / DERIVED text"],
            ["8", "Current production run / order", "Done", "Backend + Frontend", "Run bar in OC"],
            ["9", "Production history / timeline", "Done / Partial", "Backend + Frontend", "Timeline UI exists"],
            ["10", "Machine overview", "Done", "Frontend + Backend", ""],
            ["11", "Sensor center", "Done", "Frontend + Backend", ""],
            ["12", "Material profiles", "Done", "Frontend + Backend", ""],
            ["13", "Baseline manager", "Partial", "Frontend maps vs live_monitor registry", "Two systems — teams must align"],
            ["14", "Live deviations heatmap", "Done", "Frontend (RULE_BASED)", "Can later enrich from live feature evaluations"],
            ["15", "Predictions", "Remaining", "AI/ML outputs → Frontend", "Need real cards with provenance, not demo risks"],
            ["16", "Action recommendations", "Remaining", "AI/ML / rules → Frontend", "Always pair risk + action"],
            ["17", "Ticket center", "Done", "Frontend + Backend", ""],
            ["18", "Maintenance center", "Done", "Frontend + Backend", "Real RUL only when data/model exists"],
            ["19", "Energy center", "Done", "Frontend + Backend", "Needs energy source for full value"],
            ["20", "Executive view", "Done", "Frontend + Backend", ""],
        ],
    )

    doc.add_heading("3.2 Platform foundation", 2)
    add_table(
        doc,
        ["Capability", "Status", "Who"],
        [
            ["Operations Center overview API", "Done", "Backend"],
            ["Capability / locked features / setup wizard", "Partial → Done enough for demo", "Backend + Frontend"],
            ["Data source registry + provenance", "Partial", "Backend + Frontend"],
            ["Live monitor → Postgres write path", "Done", "AI/ML (BackendWriter)"],
            ["Remove SQLite from live_monitor training / debug", "Remaining", "AI/ML (+ Backend read APIs if needed)"],
            ["Validated model registry (model_versions)", "Remaining", "AI/ML → Backend → Frontend Module 5"],
            ["Event normalization / multi-machine canonical sensors", "Remaining", "Backend (+ AI/ML feature mapping)"],
            ["Quality / maintenance / material batch / energy full links", "Remaining / data", "Backend + plant integrations"],
        ],
    )

    doc.add_heading("4. What AI/ML (live_monitor) already has done", 1)
    doc.add_heading("4.1 Working today", 2)
    add_table(
        doc,
        ["Capability", "Status", "Persisted where"],
        [
            ["Poll live extruder sensors", "Done", "Backend POST /machine-raw-data/"],
            ["5-minute rolling window + features", "Done", "Backend POST /live-process-windows"],
            ["ML state classification (RandomForest, 6 states)", "Done", "Window candidate_state / confirmed_state"],
            ["Per-state Isolation Forest anomaly", "Done", "ml_anomaly_score, ml_is_anomaly, ml_model_status on run eval"],
            ["Regime baseline selection + feature z-scores", "Done", "POST /live-feature-evaluations"],
            ["Overall run evaluation (status / drift / stability)", "Done", "POST /live-run-evaluations"],
            ["Context resolve (machine / line / production run)", "Done", "From backend /machines, /production-run"],
            ["Offline manual retrain (5-min path)", "Done", "Local ml_data/*.pkl (artifacts stay on disk)"],
        ],
    )
    add_para(
        doc,
        "Machine states supported: OFF, HEATING, COOLING, READY, PRODUCTION, "
        "LOW_PRODUCTION (with confirmation windows).",
    )
    add_para(
        doc,
        "Honest product statement for AI/ML today: We deliver monitoring + state + "
        "anomaly + baseline deviation. We do not yet deliver validated scrap / RUL / "
        "quality-degradation prediction models.",
    )

    doc.add_heading("5. Remaining work — by owner", 1)
    doc.add_heading("5.A Ammar — AI/ML (live_monitor)", 2)
    add_table(
        doc,
        ["Priority", "Task", "Why", "Depends on teammate?"],
        [
            [
                "P0",
                "Postgres-only cutover — stop SQLite for raw history, retrain counting, local /live/* debug",
                "Team rule: Postgres only",
                "Yes — need read APIs / export for training windows",
            ],
            [
                "P0",
                "Stable ML payload contract for UI (fields + meaning)",
                "Teammate displays Module 7 / 15 / 16",
                "Coordinate with Backend/Frontend",
            ],
            [
                "P1",
                "Fix drift detector feature-name mismatch",
                "Architecture: where is risk?",
                "No",
            ],
            [
                "P1",
                "Align baseline-registry seeding/sync with Postgres",
                "Live z-score eval quality",
                "Backend seed/API if missing",
            ],
            [
                "P1",
                "Produce plain-language AI findings from state + anomaly + feature statuses",
                "Module 7",
                "Frontend consumes",
            ],
            [
                "P1",
                "Produce prediction / risk objects with provenance MODEL_PREDICTION or RULE_BASED + optional action",
                "Modules 15–16",
                "Frontend consumes",
            ],
            [
                "P2",
                "Fix auto-retrain to 5-minute labeled pipeline + reload state classifier in memory",
                "Ops reliability",
                "Backend history for new rows",
            ],
            [
                "P2",
                "Real data-quality fractions in FeatureEngine",
                "Guard + Data Quality story",
                "Optional Backend DQ later",
            ],
            [
                "P2",
                "PROFILE baseline path (material/profile baselines)",
                "Architecture baseline learning",
                "Backend profile IDs",
            ],
            [
                "P3",
                "Validated model_versions metrics (precision/recall/F1, false-alarm, lead time)",
                "Module 5 — only when real validation exists",
                "Backend table + Frontend display",
            ],
            [
                "Later",
                "Scrap / tool RUL / quality forecast / energy optimization models",
                "Locked until quality/maintenance/energy data",
                "Plant data + Backend connectors",
            ],
        ],
    )

    doc.add_heading("5.B Teammate — Backend", 2)
    add_table(
        doc,
        ["Priority", "Task", "Why", "Depends on Ammar?"],
        [
            ["P0", "Keep / harden live ingest APIs used by live_monitor", "Pipeline truth in Postgres", "Contract with Ammar"],
            ["P0", "Provide read APIs for training/retrain (raw history, windows, evaluations)", "Postgres-only rule", "Ammar will call them"],
            ["P1", "Ensure baseline-registry populated in Postgres", "Live evaluation quality", "Ammar defines feature list"],
            ["P1", "Optional aggregated current AI snapshot for OC", "Module 7 / 15 performance", "Uses Ammar-written tables"],
            ["P1", "Store / serve provenance fields consistently", "Architecture rule", "Ammar supplies value_source where ML"],
            ["P2", "model_versions (+ validation status) API", "Module 5", "Ammar supplies metrics after validation"],
            ["P2", "Auth on ingest if required for production", "Ops hardening", "Coordinate deploy"],
            ["P2", "Alembic completeness for all live_* tables", "Deploy reliability", "—"],
            ["P3", "Event normalization / multi-machine sensor mapping", "Scale beyond Extruder 1", "AI feature map later"],
            ["Data", "Quality, maintenance, energy, material-batch connectors", "Unlock locked features", "Not AI invent"],
        ],
    )

    doc.add_heading("5.C Teammate — Frontend", 2)
    add_para(
        doc,
        "UI shell is complete. Remaining is wiring / honesty, not redesign.",
    )
    add_table(
        doc,
        ["Priority", "Task", "Why", "Depends on Ammar?"],
        [
            ["P0", "Keep showing Prediction Readiness, never fake Accuracy", "Architecture critical rule", "No"],
            ["P1", "Module 7 — Live AI analysis: replace demo/DERIVED-only text with live eval / ML findings", "Screenshots still show placeholder AI recommendation", "Yes — Ammar must write findings"],
            ["P1", "Modules 15–16: Predictions + Actions from real API with provenance badges", "Nav exists; content still waiting", "Yes"],
            ["P1", "Mark every ML card with MODEL_PREDICTION / RULE_BASED / SIMULATED", "Trust", "Yes"],
            ["P2", "Enrich live deviations / estimated pages from live-feature-evaluations if needed", "Stronger why", "Data already from Ammar"],
            ["P2", "Remove / hide dead nav without routes, or finish routes", "Clean product", "Optional"],
            ["P2", "Wire header KI healthy to real /ai/status or model health", "Honesty", "Backend + optional Ammar model-status"],
            ["P3", "Module 5 Accuracy center UI — only when validated metrics exist", "Do not invent", "Ammar + Backend"],
            ["Done", "OC skin, map, readiness, locked features, trends, run bar, centers", "Per shared screenshots", "—"],
        ],
    )

    doc.add_heading("6. Integration contract (how teams meet)", 1)
    doc.add_heading("6.1 Ammar writes (already)", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Content"],
        [
            ["POST", "/machine-raw-data/", "Raw sensor poll"],
            ["POST", "/live-process-windows", "Features + state"],
            ["POST", "/live-run-evaluations", "Overall + ml_anomaly_* + drift/stability"],
            ["POST", "/live-feature-evaluations", "Per-feature z-score / status"],
        ],
    )
    doc.add_heading("6.2 Ammar reads (already)", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Content"],
        [
            ["GET", "/machines", "Resolve Extruder"],
            ["GET", "/production-run/ /current", "line + run context"],
            ["GET", "/baseline-registry", "Regime baselines"],
        ],
    )
    doc.add_heading("6.3 Still needed for Postgres-only AI/ML", 2)
    add_table(
        doc,
        ["Need", "Suggested owner", "Purpose"],
        [
            ["Historical raw / window export for retrain", "Backend", "Replace SQLite in build_live_windows / retrain"],
            ["Optional GET current findings feed for OC", "Backend", "Frontend Module 7 without heavy joins"],
            [
                "Agreement on finding JSON shape",
                "Both",
                "{ text, severity, value_source, state, ml_is_anomaly, feature_drivers[], recommended_action? }",
            ],
        ],
    )
    add_para(
        doc,
        "Frontend should consume for AI display: GET /live-run-evaluations, "
        "GET /live-feature-evaluations, GET /live-process-windows; later "
        "predictions/actions endpoint once defined.",
    )

    doc.add_heading("7. Simple workflow (end-to-end)", 1)
    for line in [
        "Plant sensors / Timescale",
        "        ↓",
        "Backend dashboard extruder APIs",
        "        ↓",
        "live_monitor (Ammar): buffer → features → state ML → anomaly ML → baseline eval",
        "        ↓",
        "Backend Postgres APIs (Teammate): raw + windows + feature evals + run evals",
        "        ↓",
        "Frontend Operations Center (Teammate): readiness / map / trends (done) + Live AI / Predictions / Actions (remaining wire-up)",
    ]:
        add_para(doc, line)

    doc.add_heading(
        "8. Definition of “AI/ML work complete” for current phase", 1
    )
    add_para(
        doc,
        "For the current customer-honest phase (monitor + anomaly + deviation), "
        "Ammar’s phase is complete when:",
    )
    add_bullets(
        doc,
        [
            "Live pipeline writes only to Postgres (no SQLite dependency for operation or retrain).",
            "Every cycle can produce: confirmed state, anomaly score/flag, feature deviation statuses, overall status, optional plain-language finding + optional action.",
            "Contract documented so Frontend can show Modules 7 / 14 / 15 / 16 without inventing numbers.",
            "No Accuracy % published until a validated model_versions record exists.",
        ],
    )
    add_para(
        doc,
        "Not required yet for phase complete: scrap model, tool RUL model, "
        "multi-machine trained models, energy optimization.",
    )

    doc.add_heading("9. Immediate next actions (suggested)", 1)
    doc.add_heading("9.1 Ammar (this week)", 2)
    add_bullets(
        doc,
        [
            "List every remaining SQLite touch in live_monitor and replace with backend/Postgres reads.",
            "Publish a short API payload + finding schema note for the teammate.",
            "Fix drift feature mapping.",
            "Add structured plain-language findings into run evaluation (or companion payload).",
        ],
    )
    doc.add_heading("9.2 Teammate (this week)", 2)
    add_bullets(
        doc,
        [
            "Confirm Postgres has data flowing from live_monitor (windows + run evals + feature evals).",
            "Wire OC AI recommendation / Live AI panel to real eval data (stop demo-only when live exists).",
            "Provide retrain history read API if Ammar needs it to drop SQLite.",
            "Confirm baseline-registry contents for HIGH/MID/LOW.",
        ],
    )
    doc.add_heading("9.3 Together (15–30 min sync)", 2)
    add_bullets(
        doc,
        [
            "Agree finding JSON schema.",
            "Agree when to switch OC from SIMULATED risks → MODEL_PREDICTION / RULE_BASED.",
            "Confirm evaluable states for customer UI (all states vs PRODUCTION only).",
        ],
    )

    doc.add_heading(
        "10. Checklist before claiming “prediction ready”", 1
    )
    add_bullets(
        doc,
        [
            "Quality data connected",
            "Maintenance data connected",
            "Material batches connected",
            "Fault / scrap labels linked to runs",
            "Model validated and stored in model_versions",
            "UI shows real metrics under Module 5 — not readiness renamed as Accuracy",
        ],
    )
    add_para(
        doc,
        "Until then: monitor + readiness + locked features is the honest product story.",
    )
    add_para(
        doc,
        "Document owner: Ammar (AI/ML) + Backend/Frontend teammate — keep this "
        "file updated when ownership or status changes.",
    )

    out = Path(__file__).with_name("WORK_ALIGNMENT_ARCHITECTURE_vs_TEAMS.docx")
    doc.save(out)
    print(f"saved {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
